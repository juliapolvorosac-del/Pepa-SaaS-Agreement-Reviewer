import streamlit as st
import anthropic
import io
from pathlib import Path

import docx as python_docx
import pdfplumber


# =====================
# --- CSS y tema ---
# =====================

CUSTOM_CSS = """
<style>
    /* Ocultar marca de Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    header[data-testid="stHeader"] {background: transparent;}

    /* Cabecera principal PEPA */
    .pepa-header {
        background: linear-gradient(135deg, #1B3A6B 0%, #2D5A9E 100%);
        padding: 28px 36px;
        border-radius: 12px;
        margin-bottom: 28px;
        box-shadow: 0 4px 16px rgba(27,58,107,0.15);
    }
    .pepa-header h1 {
        color: #FFFFFF;
        margin: 0;
        font-size: 32px;
        font-weight: 700;
        letter-spacing: -0.5px;
    }
    .badge-beta {
        display: inline-block;
        background-color: #C9A84C;
        color: #1B3A6B;
        font-size: 11px;
        font-weight: 800;
        letter-spacing: 1.5px;
        padding: 3px 9px;
        border-radius: 20px;
        vertical-align: middle;
        margin-left: 12px;
        text-transform: uppercase;
        position: relative;
        top: -4px;
    }
    .pepa-header p {
        color: #C9A84C;
        margin: 6px 0 0 0;
        font-size: 15px;
        font-weight: 400;
    }
    .pepa-header .aviso-cabecera {
        color: #FFE0A0;
        margin: 12px 0 0 0;
        font-size: 13px;
        font-weight: 400;
        border-top: 1px solid rgba(255,255,255,0.2);
        padding-top: 10px;
    }

    /* Botón principal */
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #1B3A6B, #2D5A9E);
        border: none;
        color: white;
        font-size: 16px;
        font-weight: 600;
        padding: 14px 24px;
        border-radius: 8px;
        transition: opacity 0.2s;
    }
    .stButton > button[kind="primary"]:hover {
        opacity: 0.88;
    }

    /* Informe de revisión */
    .informe-container {
        background: #F4F6F9;
        border-left: 4px solid #1B3A6B;
        border-radius: 0 8px 8px 0;
        padding: 24px 28px;
        margin-top: 8px;
        color: #444444;
    }
    .informe-container h1, .informe-container h2,
    .informe-container h3, .informe-container h4 {
        color: #1B3A6B;
    }
    .informe-container strong {
        color: #1B2340;
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #E8E8E8;
        border-right: 2px solid #CCCCCC;
    }
    [data-testid="stSidebar"] h2 {
        color: #333333;
        font-size: 14px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] li,
    [data-testid="stSidebar"] .stMarkdown {
        color: #444444;
    }
    [data-testid="stSidebar"] .stExpander {
        background-color: #DCDCDC;
        border: 1px solid #BBBBBB;
        border-radius: 6px;
    }
    [data-testid="stSidebar"] hr {
        border-color: #BBBBBB;
    }

    /* Divisor */
    hr {
        border-color: #E0E4EC;
    }

    /* Aviso legal en sidebar */
    .aviso-legal {
        font-size: 11px;
        color: #888;
        border-top: 1px solid #E0E4EC;
        padding-top: 12px;
        margin-top: 12px;
        line-height: 1.5;
    }
</style>
"""


# --- Helpers: extracción de texto ---

def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".txt"):
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    if name.endswith(".docx"):
        doc = python_docx.Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if name.endswith(".pdf"):
        pages = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)

    return ""


def extract_text_from_path(path: Path) -> str:
    name = path.name.lower()

    if name.endswith(".txt"):
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return path.read_text(encoding=encoding).strip()
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="replace").strip()

    if name.endswith(".docx"):
        doc = python_docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if name.endswith(".pdf"):
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)

    return ""


# --- Helpers: archivos de referencia ---

BASE_DIR = Path(__file__).parent


def load_instructions() -> str:
    path = BASE_DIR / "instrucciones.txt"
    if path.exists():
        return extract_text_from_path(path)
    return ""


def load_manual() -> tuple[str, str]:
    for candidate in ("manual.txt", "manual.docx", "manual.pdf"):
        path = BASE_DIR / candidate
        if path.exists():
            return extract_text_from_path(path), candidate
    return "", ""


def get_api_key() -> str:
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        return ""


# --- Análisis con Claude ---

def stream_report(contract_text: str, instructions: str, manual_text: str, api_key: str):
    client = anthropic.Anthropic(api_key=api_key)

    manual_block = ""
    if manual_text.strip():
        manual_block = f"""---

MANUAL DE REFERENCIA DE LA EMPRESA (usa este documento para comparar y clasificar cada cláusula):
{manual_text}

"""

    prompt = f"""INSTRUCCIONES DE REVISIÓN:
{instructions}

{manual_block}---

CONTRATO A REVISAR:
{contract_text}

---

Genera el informe de revisión siguiendo exactamente las instrucciones proporcionadas.
Usa formato Markdown con secciones bien definidas, negritas para conceptos clave y listas cuando proceda.
Escribe en español."""

    with client.messages.stream(
        model="claude-opus-4-5",
        max_tokens=8096,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        for text in stream.text_stream:
            yield text


# =====================
# --- Interfaz ---
# =====================

st.set_page_config(
    page_title="PEPA · Revisor de Contratos",
    page_icon="⚖️",
    layout="wide",
)

# Inyectar CSS personalizado
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# Cabecera PEPA
st.markdown("""
<div class="pepa-header">
    <h1>⚖️ PEPA — Revisión preliminar de contratos SaaS <span class="badge-beta">Beta</span></h1>
    <p>Plataforma de revisión contractual preliminar con IA de conformidad con los estándares de PEPA</p>
    <p class="aviso-cabecera">⚠️ El presente proyecto es educativo y no sustituye el asesoramiento jurídico de un profesional cualificado</p>
</div>
""", unsafe_allow_html=True)

# --- Barra lateral ---
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/scales.png", width=56)
    st.markdown("## PEPA Legal Tech")
    st.markdown("Revisión automatizada de contratos SaaS basada en el manual de posiciones jurídicas de PEPA.")
    st.divider()

    st.markdown("## Cómo usar")
    st.markdown("""
1. Sube el contrato (`.txt`, `.docx` o `.pdf`).
2. Pulsa **Generar informe**.
3. Revisa el análisis por cláusulas.
4. Descarga el informe si lo necesitas.
""")
    st.divider()

    st.markdown("## Estado del sistema")

    instructions_preview = load_instructions()
    manual_text_sidebar, manual_name_sidebar = load_manual()

    if instructions_preview:
        st.success("✅ Instrucciones cargadas")
        with st.expander("Ver instrucciones"):
            st.text(instructions_preview[:600] + ("..." if len(instructions_preview) > 600 else ""))
    else:
        st.error("❌ instrucciones.txt no encontrado")

    if manual_text_sidebar:
        st.success(f"✅ Manual cargado ({len(manual_text_sidebar):,} car.)")
        with st.expander("Ver extracto del manual"):
            st.text(manual_text_sidebar[:600] + ("..." if len(manual_text_sidebar) > 600 else ""))
    else:
        st.warning("⚠️ Manual no encontrado")

    st.markdown("""
<div class="aviso-legal">
⚠️ Este análisis es orientativo y no sustituye el asesoramiento jurídico de un profesional cualificado.
</div>
""", unsafe_allow_html=True)

# --- Comprobación de API key ---
api_key = get_api_key()
if not api_key:
    st.error(
        "❌ No se encontró la API key de Google. "
        "Añádela en **Settings → Secrets** de tu app en Streamlit Cloud."
    )
    st.stop()

# --- Subida de archivo ---
st.markdown("### 📂 Documento a revisar")
uploaded_file = st.file_uploader(
    "Sube el contrato",
    type=["txt", "docx", "pdf"],
    label_visibility="collapsed",
)

if uploaded_file is None:
    st.info("👆 Sube un contrato en formato .txt, .docx o .pdf para comenzar.")
    st.stop()

# --- Extracción de texto ---
with st.spinner("Leyendo documento..."):
    contract_text = extract_text(uploaded_file)

if not contract_text.strip():
    st.error("No se pudo extraer texto del documento. Verifica que no esté vacío ni protegido con contraseña.")
    st.stop()

col1, col2 = st.columns([3, 1])
with col1:
    st.success(f"✅ **{uploaded_file.name}** cargado correctamente ({len(contract_text):,} caracteres)")
with col2:
    with st.expander("Ver texto extraído"):
        st.text_area(
            "Contenido",
            contract_text[:4000] + ("\n\n[...]" if len(contract_text) > 4000 else ""),
            height=250,
            disabled=True,
            label_visibility="collapsed",
        )

# --- Botón de análisis ---
st.markdown("---")
instructions = load_instructions()
manual_text, manual_name = load_manual()

if not instructions:
    st.warning("⚠️ El archivo `instrucciones.txt` está vacío.")
    st.stop()

if manual_text:
    st.info(f"📘 Manual de referencia: **{manual_name}** ({len(manual_text):,} caracteres)")
else:
    st.warning("⚠️ Manual no encontrado — la revisión usará estándares genéricos de mercado.")

if st.button("🔍 Generar informe de revisión", type="primary", use_container_width=True):
    st.markdown("---")
    st.markdown("### 📊 Informe de Revisión")

    report_area = st.empty()
    full_report = ""

    try:
        for chunk in stream_report(contract_text, instructions, manual_text, api_key):
            full_report += chunk
            report_area.markdown(full_report + "▌")
        report_area.markdown(full_report)

    except anthropic.AuthenticationError:
        st.error("❌ API key incorrecta. Revisa el valor en los Secrets de Streamlit Cloud.")
        st.stop()
    except anthropic.RateLimitError:
        st.error("❌ Límite de uso alcanzado. Espera unos minutos e inténtalo de nuevo.")
        st.stop()
    except Exception as e:
        st.error(f"❌ Error inesperado: {e}")
        st.stop()

    st.markdown("---")
    st.download_button(
        label="⬇️ Descargar informe (.md)",
        data=full_report,
        file_name=f"informe_{Path(uploaded_file.name).stem}.md",
        mime="text/markdown",
        use_container_width=True,
    )
